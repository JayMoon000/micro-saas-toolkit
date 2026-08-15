import os
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

# 저장 경로 설정
target_dir = r"D:\Gemini_Files\Micro_SaaS_Toolkit"
os.makedirs(target_dir, exist_ok=True)

# ----------------------------------------------------
# 1. Micro_SaaS_Toolkit_Progress_Report.docx 생성
# ----------------------------------------------------
doc = docx.Document()

# 기본 폰트 설정 (맑은 고딕)
style = doc.styles['Normal']
font = style.font
font.name = '맑은 고딕'
font.size = Pt(10.5)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)

# 테이블 셀 배경색 설정 함수
def set_cell_background(cell, hex_color):
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

# 제목 추가
title_p = doc.add_paragraph()
title_run = title_p.add_run("Micro SaaS & 서브도메인 툴킷 프로젝트 개발 진행 보고서")
title_run.font.size = Pt(18)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
title_p.paragraph_format.space_after = Pt(12)

# 개요
h1 = doc.add_heading(level=1)
h1_run = h1.add_run("1. 프로젝트 개요")
h1_run.font.name = '맑은 고딕'
h1_run.font.size = Pt(14)
h1_run.font.bold = True
h1_run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)

p1 = doc.add_paragraph("본 프로젝트는 서버비 $0 기반의 초슬림 유틸리티 툴킷(tool.starsign16.com)을 개발하고, 트래픽 데이터 기반의 검증을 거쳐 모바일 앱 및 생산성 플랫폼으로 확장하는 것을 목표로 합니다.")
p1.paragraph_format.space_after = Pt(10)

# 개발 완료 및 진행 현황
h2 = doc.add_heading(level=1)
h2_run = h2.add_run("2. 툴킷 개발 현황")
h2_run.font.name = '맑은 고딕'
h2_run.font.size = Pt(14)
h2_run.font.bold = True
h2_run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)

table = doc.add_table(rows=5, cols=4)
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ["구분", "파일명", "주요 구현 기능", "상태"]
data = [
    ["Tool 1", "flowchart_maker.html", "Visual Builder(노코드 블록) 방식 텍스트 플로우차트 생성기 (Mermaid.js 연동)", "완료"],
    ["Tool 2", "epub_maker.html", "5개 포맷(PDF, DOCX, MD, HTML, TXT) 지원 1초 EPUB 전자책 변환기 (JSZip, pdf.js, mammoth.js)", "완료"],
    ["Tool 3", "fast_calendar.html", "자연어 메모/약속 파싱 구글 캘린더 원클릭 퀵 등록 툴 (정규식 기반 연동)", "완료"],
    ["Tool 4", "ocr_excel_maker.html", "영수증/표 OCR 스캔 및 엑셀 셀 편집기 (Tesseract.js 한계로 보류)", "보류"]
]

# 헤더 스타일링
hdr_cells = table.rows[0].cells
for idx, text in enumerate(headers):
    hdr_cells[idx].text = text
    set_cell_background(hdr_cells[idx], "2563EB")
    p = hdr_cells[idx].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs:
        r.font.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

# 데이터 셀 입력
for row_idx, row_data in enumerate(data):
    row_cells = table.rows[row_idx + 1].cells
    for col_idx, cell_value in enumerate(row_data):
        row_cells[col_idx].text = cell_value
        p = row_cells[col_idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx in [0, 1, 3] else WD_ALIGN_PARAGRAPH.LEFT
        if row_idx % 2 == 1:
            set_cell_background(row_cells[col_idx], "F8FAFC")

# 향후 계획
h3 = doc.add_heading(level=1)
h3_run = h3.add_run("3. 향후 계획")
h3_run.font.name = '맑은 고딕'
h3_run.font.size = Pt(14)
h3_run.font.bold = True
h3_run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)

p2 = doc.add_paragraph("• 3.5 AI 블로그용 요약 인포그래픽/차트 카드 생성기 개발\n• 보류된 OCR 툴의 Cloudflare Workers AI 비전 모델 전환 검토\n• 메인 도메인 애드센스 승인 완료 후 서브도메인(tool.starsign16.com) 정식 연결")
p2.paragraph_format.space_after = Pt(10)

docx_path = os.path.join(target_dir, "Micro_SaaS_Toolkit_Progress_Report.docx")
doc.save(docx_path)


# ----------------------------------------------------
# 2. Micro_SaaS_Toolkit_Handover.xlsx 생성
# ----------------------------------------------------
wb = openpyxl.Workbook()
ws = wb.active
ws.title = "핸드오버 요약표"

# 헤더 작성
excel_headers = ["구분", "아이템/파일명", "주요 구현 기능 및 특이사항", "상태", "비고"]
ws.append(excel_headers)

# 데이터 작성
excel_data = [
    ["Tool 1", "flowchart_maker.html", "사각형/마름모 노코드 블록 입력, 스마트 연결 추적, SVG 다운로드", "완료", "100% Client-Side"],
    ["Tool 2", "epub_maker.html", "PDF, DOCX, MD, HTML, TXT 드래그 앤 드롭 및 1초 EPUB 압축 생성", "완료", "pdf.js, JSZip"],
    ["Tool 3", "fast_calendar.html", "날짜/시간/장소 자연어 파싱, 구글 캘린더 URL 자동 세팅", "완료", "Regex 파서 적용"],
    ["Tool 4", "ocr_excel_maker.html", "영수증/표 OCR 스캔 및 엑셀 셀 편집기", "보류", "Workers AI 적용 검토"],
    ["Next", "chart_card_maker.html", "AI 블로그용 요약 인포그래픽/차트 카드 생성기", "대기", "다음 세션 예정"]
]

for row in excel_data:
    ws.append(row)

# 스타일 적용
header_fill = PatternFill(start_color="2563EB", end_color="2563EB", fill_type="solid")
header_font = Font(name="맑은 고딕", size=11, bold=True, color="FFFFFF")
data_font = Font(name="맑은 고딕", size=10)
align_center = Alignment(horizontal="center", vertical="center")
align_left = Alignment(horizontal="left", vertical="center")

thin_border = Border(
    left=Side(style='thin', color='CBD5E1'),
    right=Side(style='thin', color='CBD5E1'),
    top=Side(style='thin', color='CBD5E1'),
    bottom=Side(style='thin', color='CBD5E1')
)

# 헤더 스타일링
for col in range(1, len(excel_headers) + 1):
    cell = ws.cell(row=1, column=col)
    cell.fill = header_fill
    cell.font = header_font
    cell.alignment = align_center
    cell.border = thin_border

# 데이터 셀 스타일링
for row in range(2, len(excel_data) + 2):
    for col in range(1, len(excel_headers) + 1):
        cell = ws.cell(row=row, column=col)
        cell.font = data_font
        cell.border = thin_border
        cell.alignment = align_center if col in [1, 2, 4, 5] else align_left

# 열 너비 자동 조절
for col in ws.columns:
    max_len = 0
    col_letter = get_column_letter(col[0].column)
    for cell in col:
        if cell.value:
            # 한글 글자 수 판별 포함 개략적 너비 계산
            val_str = str(cell.value)
            korean_count = sum(1 for c in val_str if ord(c) > 128)
            length = len(val_str) + korean_count
            if length > max_len:
                max_len = length
    ws.column_dimensions[col_letter].width = max(max_len + 3, 12)

xlsx_path = os.path.join(target_dir, "Micro_SaaS_Toolkit_Handover.xlsx")
wb.save(xlsx_path)

print(f"파일 생성 완료:\n- {docx_path}\n- {xlsx_path}")
코드생성끝