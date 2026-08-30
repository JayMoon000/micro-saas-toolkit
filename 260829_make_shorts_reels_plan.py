import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

SAVE_DIR = r"D:\Gemini_Files\Micro_SaaS_Toolkit"
FILE_NAME = "Shorts_Reels_Video_Script_BackgroundRemover.docx"

def set_cell_background(cell, fill_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
    tcPr.append(shd)

def create_script_docx():
    os.makedirs(SAVE_DIR, exist_ok=True)
    doc = Document()

    # 페이지 여백
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # 타이틀
    title_p = doc.add_paragraph()
    title_run = title_p.add_run("🎬 [숏폼/릴스] StarSign16 AI 누끼 제거 툴 홍보 영상 콘티")
    title_run.font.name = "Malgun Gothic"
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(15, 23, 42)
    title_p.paragraph_format.space_after = Pt(12)

    # 개요 메타정보
    meta_p = doc.add_paragraph()
    meta_p.add_run("• 타깃: ").bold = True
    meta_p.add_run("디자이너, 마케터, 블로거, 쇼핑몰 셀러, 직장인\n")
    meta_p.add_run("• 핵심 셀링포인트: ").bold = True
    meta_p.add_run("100% 무료, 로그인 없음, 워터마크 없음, 서버 업로드 없는 안전한 로컬 AI 연산\n")
    meta_p.add_run("• 영상 길이: ").bold = True
    meta_p.add_run("30초 내외 (9:16 세로형 쇼츠/릴스 규격)\n")
    meta_p.add_run("• 목표 링크: ").bold = True
    meta_p.add_run("https://tool.starsign16.com/ko/background-remover")
    meta_p.paragraph_format.space_after = Pt(16)

    # 표 생성 (영상 타임라인 콘티)
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    headers = ["시간 (초)", "화면 연출 (Video)", "오디오 / 내레이션 (Audio)", "자막 키워드 (Text)"]
    hdr_cells = table.rows[0].cells
    widths = [Inches(1.0), Inches(2.3), Inches(2.3), Inches(1.4)]

    for idx, name in enumerate(headers):
        hdr_cells[idx].text = name
        hdr_cells[idx].width = widths[idx]
        set_cell_background(hdr_cells[idx], "1E293B")
        p = hdr_cells[idx].paragraphs[0]
        p.runs[0].font.name = "Malgun Gothic"
        p.runs[0].font.bold = True
        p.runs[0].font.color.rgb = RGBColor(255, 255, 255)
        p.runs[0].font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    data = [
        ("00:00 ~ 00:03\n(3초 훅)", 
         "유료 결제 팝업창이나 회원가입 창 보며 답답해하는 제스처 / 화면 캡처", 
         "아직도 사진 누끼 따려고 유료 결제하거나 회원가입 하세요?", 
         "🚨 아직도 누끼 딸 때\n결제하세요?"),
        
        ("00:03 ~ 00:08\n(Pain Point)", 
         "기존 무료 툴 이용 시 워터마크 찍히거나 하루 1장 제한 걸리는 화면", 
         "포토샵 켜기는 귀찮고, 무료 툴은 화질 깨지거나 워터마크 박히잖아요.", 
         "❌ 워터마크\n❌ 하루 1장 제한\n❌ 회원가입"),
        
        ("00:08 ~ 00:15\n(해결책 제시)", 
         "브라우저에서 tool.starsign16.com 접속 후 복잡한 배경 사진을 드래그앤드롭하는 실제 화면", 
         "그냥 브라우저 켜고 사진을 툭 던져 넣으세요.", 
         "✨ 설치 NO / 로그인 NO\n그냥 드래그앤드롭!"),
         
        ("00:15 ~ 00:24\n(기능 시연)", 
         "부드러운 로딩 게이지 작동 후 피사체만 완벽하게 분리되어 투명 바둑판 배경으로 바뀌는 결과 화면 확대", 
         "서버로 이미지를 보내지 않고 내 컴퓨터에서 AI가 직접 땁니다. 털 한 올까지 깔끔하게 분리되죠?", 
         "🔒 서버 전송 ZERO\n고정밀 AI 로컬 연산"),
         
        ("00:24 ~ 00:30\n(Call To Action)", 
         "투명 PNG 다운로드 버튼 클릭 ➔ 원본 화질 파일 다운 완료 ➔ 툴킷 메인 허브 화면", 
         "워터마크 없는 100% 무료 누끼 툴. 고정 댓글 링크에서 바로 써보세요!", 
         "👉 고정 댓글 링크\nStarSign16 툴킷")
    ]

    for row_idx, item in enumerate(data):
        row_cells = table.add_row().cells
        bg_color = "F8FAFC" if row_idx % 2 == 0 else "FFFFFF"
        for col_idx, text in enumerate(item):
            row_cells[col_idx].text = text
            row_cells[col_idx].width = widths[col_idx]
            set_cell_background(row_cells[col_idx], bg_color)
            p = row_cells[col_idx].paragraphs[0]
            p.runs[0].font.name = "Malgun Gothic"
            p.runs[0].font.size = Pt(9.5)
            if col_idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 하단 캡션 가이드
    doc.add_paragraph().paragraph_format.space_before = Pt(16)
    p_caption = doc.add_paragraph()
    p_caption.add_run("📌 유튜브 쇼츠 / 릴스 본문 & 고정 댓글 텍스트\n").bold = True
    caption_body = (
        "[본문 캡션]\n"
        "회원가입, 결제 유도, 워터마크 전부 없는 100% 무료 AI 누끼 제거 툴 ✂️\n"
        "내 컴퓨터 브라우저 메모리에서 직접 연산되어 서버로 사진이 전송되지 않아 보안도 안전합니다.\n\n"
        "👉 사용 링크: tool.starsign16.com/ko/background-remover\n\n"
        "#누끼따기 #AI툴 #무료누끼 #생산성도구 #웹유틸리티 #StarSign16\n\n"
        "[고정 댓글]\n"
        "로그인 없이 3초 만에 누끼 따기 👉 https://tool.starsign16.com/ko/background-remover"
    )
    p_caption.add_run(caption_body)

    file_path = os.path.join(SAVE_DIR, FILE_NAME)
    doc.save(file_path)
    print(f"📄 Generated docx: {file_path}")

if __name__ == "__main__":
    create_script_docx()