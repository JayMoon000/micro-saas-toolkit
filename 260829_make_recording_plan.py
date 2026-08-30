import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

SAVE_DIR = r"D:\Gemini_Files\Micro_SaaS_Toolkit"
FILE_NAME = "Shorts_Screen_Recording_Guide.docx"

def set_cell_background(cell, fill_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
    tcPr.append(shd)

def create_recording_guide():
    os.makedirs(SAVE_DIR, exist_ok=True)
    doc = Document()

    # 여백 설정
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # 타이틀
    title_p = doc.add_paragraph()
    title_run = title_p.add_run("📹 StarSign16 AI 누끼 제거 숏폼 화면 녹화 실전 가이드")
    title_run.font.name = "Malgun Gothic"
    title_run.font.size = Pt(17)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(15, 23, 42)
    title_p.paragraph_format.space_after = Pt(12)

    # 1. 사전 녹화 환경 세팅
    h1 = doc.add_paragraph()
    h1.add_run("1. 사전 녹화 환경 세팅").bold = True
    h1.runs[0].font.size = Pt(13)
    h1.runs[0].font.name = "Malgun Gothic"
    
    settings = [
        ("화면 비율", "1080 x 1920 (9:16 세로형) 또는 브라우저 창 폭을 480~520px로 좁혀 세로 녹화"),
        ("브라우저 설정", "북마크바 숨기기 (Ctrl+Shift+B), 크롬 배율 110~125% 확대 (폰 화면 가독성 확보)"),
        ("타깃 URL", "https://tool.starsign16.com/ko/background-remover"),
        ("준비 샘플 이미지", "배경과 인물/사물이 확실히 구분되면서도 머리카락이나 윤곽선이 돋보이는 고화질 사진 1장"),
        ("녹화 도구", "OBS Studio (세로 캔버스 1080x1920 세팅) 또는 윈도우 캡처(Win+Alt+R) 후 편집기에서 크롭")
    ]
    for k, v in settings:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"• {k}: ").bold = True
        p.add_run(v)
        p.paragraph_format.space_after = Pt(3)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # 2. 클립별 녹화 액션 시퀀스
    h2 = doc.add_paragraph()
    h2.add_run("2. 클립별 녹화 액션 시퀀스 (5개 컷)").bold = True
    h2.runs[0].font.size = Pt(13)
    h2.runs[0].font.name = "Malgun Gothic"

    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    headers = ["컷 번호", "녹화 화면 및 대상", "마우스 / 사용자 액션", "편집 시 강조 포인트"]
    hdr_cells = table.rows[0].cells
    widths = [Inches(0.9), Inches(2.2), Inches(2.3), Inches(1.8)]

    for idx, name in enumerate(headers):
        hdr_cells[idx].text = name
        hdr_cells[idx].width = widths[idx]
        set_cell_background(hdr_cells[idx], "0F172A")
        p = hdr_cells[idx].paragraphs[0]
        p.runs[0].font.name = "Malgun Gothic"
        p.runs[0].font.bold = True
        p.runs[0].font.color.rgb = RGBColor(255, 255, 255)
        p.runs[0].font.size = Pt(9.5)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    data = [
        ("Cut 1\n(0~4초)", 
         "타사 유료 사이트 or 결제 유도 팝업창", 
         "유료 결제창 보며 마우스로 닫기 누르거나 망설이는 모션", 
         "붉은색 X 표시 또는 '아직도 돈 내고 씀?' 자막 타격"),
        
        ("Cut 2\n(4~10초)", 
         "StarSign16 누끼 페이지\n(다크 UI)", 
         "탐색기에서 샘플 사진을 드래그하여 중앙 박스에 부드럽게 Drop", 
         "마우스 드래그 궤적 줌인 + '회원가입 없음' 강조"),
        
        ("Cut 3\n(10~18초)", 
         "로딩 오버레이 화면", 
         "부드러운 스피너와 'AI 모델 메모리 로드 중...' 텍스트 자연스럽게 노출 (2~3초 분량만 편집 컷)", 
         "배속(2x~3x) 처리하여 지루함 없애고 🔒 100% 온디바이스 배지 확대"),
        
        ("Cut 4\n(18~24초)", 
         "원본 vs 투명 누끼 결과창\n(체커보드 배경)", 
         "마우스로 투명 배경 결과물 주변을 가리키며 깔끔한 엣지 확인", 
         "체커보드 투명 바둑판 부분 확대 (화질 저하 없음 강조)"),
        
        ("Cut 5\n(24~30초)", 
         "다운로드 및 메인 허브", 
         "'투명 PNG 다운로드' 클릭 ➔ 좌측 상단 '← StarSign16 툴킷 허브' 클릭", 
         "다운로드 즉시 완료 화면 + '무료 툴 16종 허브' 자막 및 댓글 유도")
    ]

    for row_idx, item in enumerate(data):
        row_cells = table.add_row().cells
        bg_color = "F8FAFC" if row_idx % 2 == 0 else "FFFFFF"
        for col_idx, text in enumerate(item):
            row_cells[col_idx].text = text
            row_cells[col_idx].width = widths[col_idx]
            set_cell_background(row_cells[col_idx], bg_color)
            p = row_cells[col_idx].paragraphs[0]
            p.runs[0].font.name = "Malgun Gothic"
            p.runs[0].font.size = Pt(9)
            if col_idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 3. 편집 팁
    doc.add_paragraph().paragraph_format.space_before = Pt(12)
    h3 = doc.add_paragraph()
    h3.add_run("3. 캡컷(CapCut) / 브루(Vrew) 편집 핵심 팁").bold = True
    h3.runs[0].font.size = Pt(13)
    h3.runs[0].font.name = "Malgun Gothic"

    tips = [
        ("로딩 구간 배속", "실제 25초 연산 중 로딩창 화면은 편집기에서 2~3초로 압축(3~4배속)하여 빠른 템포 유지"),
        ("BGM 선정", "비트감 있고 경쾌한 로열티 프리 테크/생산성 숏폼 음원 배치"),
        ("자막 스타일", "중앙 하단에 검은색 배경 박스가 들어간 노란색/흰색 볼드 폰트로 1~2줄씩 전환"),
        ("CTA 배치", "마지막 3초 동안 '댓글 링크 클릭' 화살표 스티커 깜빡임 효과 적용")
    ]
    for k, v in tips:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"• {k}: ").bold = True
        p.add_run(v)
        p.paragraph_format.space_after = Pt(3)

    file_path = os.path.join(SAVE_DIR, FILE_NAME)
    doc.save(file_path)
    print(f"📄 Generated Guide docx: {file_path}")

if __name__ == "__main__":
    create_recording_guide()